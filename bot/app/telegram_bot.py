import telebot
import datetime
import psycopg2
from contextlib import closing
from config import TOKEN, PG_HOST, PG_NAME, PG_USER, PG_PASSWORD
from datetime import datetime as dtm
import requests
from docx import Document
from docx.shared import Pt
import os
from words_cloud import get_words_cloud_picture



bot = telebot.TeleBot(TOKEN)
UserRequests = {}

@bot.message_handler(commands=['start'])
def start_message(message):
    keyboard = telebot.types.ReplyKeyboardMarkup(True)
    keyboard.row('Облако слов', 'Статистика')
    #keyboard.row('Негативные комментарии', 'Позитивные комментарии')
    keyboard.row('Комментарии по рейтингу', 'Комментарии по тональности')
    bot.send_message(message.chat.id, 'Приветствую! Этот бот может помочь проанализировать отзывы на продукты бренда. Выберете кнопку из меню для того, чтобы получить комментарии в удобном для вас формате', reply_markup=keyboard)


@bot.message_handler(commands=['info'])
def start_message(message):
    keyboard = telebot.types.ReplyKeyboardMarkup(True)
    keyboard.row('Облако слов', 'Статистика')
    #keyboard.row('Негативные комментарии', 'Позитивные комментарии')
    keyboard.row('Комментарии по рейтингу', 'Комментарии по тональности')
    #keyboard.row('Выдача комментариев по рейтингу', 'Выдача комментариев по тональности')
    bot.send_message(message.chat.id, 'Этот бот может помочь проанализировать отзывы на продукты бренда. Выберете кнопку из меню для того, чтобы получить комментарии в удобном для вас формате', reply_markup=keyboard)


@bot.message_handler(commands=['help'])
def help_message(message):
    bot.send_message(message.chat.id, 'Добрый день! Этот бот может помочь проанализировать отзывы на продукты.')


@bot.message_handler(content_types=['text'])
def send_text(message):
    chat_id = message.chat.id
    if message.text.lower() == 'статистика':
        #message.text.lower() == 'негативные комментарии' or message.text.lower() == 'позитивные комментарии':
        request_category = message.text
        UserRequests[chat_id] = {"id": chat_id, "request_category": request_category,
                                 "platform": None, "start_date": None,"end_date": None, "brand": None}
        platform_choosing(message)

    elif message.text.lower() == 'комментарии по рейтингу' or \
            message.text.lower() == 'комментарии по тональности':
        if message.text.lower() == 'комментарии по рейтингу':
            bot.send_message(message.chat.id, text="Позитивные комментарии имеют рейтинг 4 или 5, в то время как негативные - 3 и ниже.")
        request_category = message.text
        UserRequests[chat_id] = {"id": chat_id, "request_category": request_category,
                                 "platform": None, "start_date": None, "end_date": None, "brand": None}
        markup = telebot.types.InlineKeyboardMarkup()
        #markup.add(telebot.types.InlineKeyboardButton(text='Общее облако слов', callback_data='Общее облако слов'))
        markup.add(telebot.types.InlineKeyboardButton(text='Позитивные комментарии',
                                                      callback_data='Позитивные комментарии'))
        markup.add(telebot.types.InlineKeyboardButton(text='Негативные комментарии',
                                                      callback_data='Негативные комментарии'))
        bot.send_message(message.chat.id, text="Выберите, какие комментарии вы хотите получить",
                         reply_markup=markup)

    elif message.text.lower() == 'облако слов':
        markup = telebot.types.InlineKeyboardMarkup()
        markup.add(telebot.types.InlineKeyboardButton(text='Общее облако слов', callback_data='Общее облако слов'))
        markup.add(telebot.types.InlineKeyboardButton(text='Облако позитивных комментариев', callback_data='Облако позитивных комментариев'))
        markup.add(telebot.types.InlineKeyboardButton(text='Облако негативных комментариев', callback_data='Облако негативных комментариев'))
        bot.send_message(message.chat.id, text="Выберите, из каких комментариев строить облако слов", reply_markup=markup)





@bot.message_handler(content_types = ['text'])
def platform_choosing(message):
    markup = telebot.types.InlineKeyboardMarkup()
    markup.add(telebot.types.InlineKeyboardButton(text='Wildberries', callback_data='Wildberries'))
    markup.add(telebot.types.InlineKeyboardButton(text='Pet Shop', callback_data='Pet Shop'))
    markup.add(telebot.types.InlineKeyboardButton(text='Мир корма', callback_data='Мир корма'))
    bot.send_message(message.chat.id, text="Выберете платформу", reply_markup=markup)


def get_standart_date_interval(weeks_num):
    now = datetime.datetime.now()
    start_date = str(now.date() - datetime.timedelta(weeks=weeks_num))
    end_date = str(now.date())
    return start_date, end_date


def check_date(date):
    date_format = "%Y-%m-%d"
    try:
        res = bool(dtm.strptime(date, date_format))
    except ValueError:
        res = False
    return res


def dates_interval_choosing(message):
    start_date = bot.send_message(message.chat.id, "Введите начальную дату в формате '2023-01-20'")
    bot.register_next_step_handler(message, start_date_check)


def start_date_check(message):
    start_date = message.text
    res = check_date(start_date)
    if res:
        dates_end_choosing(message, start_date)
    else:
        error = bot.send_message(message.chat.id, "Неверный формат")
        dates_interval_choosing(message)


def dates_end_choosing(message, start_date):
    end_date = bot.send_message(message.chat.id, "Введите конечную дату в формате '2023-01-20'")
    bot.register_next_step_handler(message, dates_prepare,start_date)


def dates_prepare(message, start_date):
    end_date = message.text
    res = check_date(end_date)
    if res:
        if start_date <= end_date:
            dates_interval_set(message, start_date, end_date)
        else:
            error = bot.send_message(message.chat.id, "Дата начала больше конечной даты, попробуйте заново. " +
                                    "Например, допустимый промежуток - от 2023-01-01 до 2023-01-05.")
            dates_end_choosing(message, start_date)
    else:
        error = bot.send_message(message.chat.id, "Неверный формат")
        dates_end_choosing(message, start_date)




def dates_interval_set(message, start_date, end_date):
    try:
        UserRequests[message.chat.id]['start_date'] = start_date
        UserRequests[message.chat.id]['end_date'] = end_date
        brand_choosing(message)
    except:
        bot.send_message(message.chat.id, 'Пожалуйста, начните с начала и выберите категорию из меню')


def brand_choosing(message):
    brand = bot.send_message(message.chat.id, 'Введите бренд (например, Whiskas)')
    if str(UserRequests[message.chat.id]['request_category']).lower() == 'комментарии по рейтингу' or \
            str(UserRequests[message.chat.id]['request_category']).lower() == 'комментарии по тональности':
        bot.register_next_step_handler(message, choose_format)
    else:
        bot.register_next_step_handler(message, set_brand)


def choose_format(message):
    UserRequests[message.chat.id]['brand'] = message.text
    markup = telebot.types.InlineKeyboardMarkup()
    markup.add(telebot.types.InlineKeyboardButton(text='Word', callback_data='Word'))
    markup.add(telebot.types.InlineKeyboardButton(text='JSON', callback_data='JSON'))
    bot.send_message(message.chat.id, text="Выберете, в каком формате вам удобно получить комментарии",
                     reply_markup=markup)


def set_brand(message):
    chat_id = message.chat.id
    brand = message.text
    try:
        UserRequests[chat_id]['brand'] = brand
        check_info(message)
    except Exception as e:
        bot.send_message(message.chat.id, 'Пожалуйста, начните с начала и выберите категорию из меню')



def check_info(message):
    chat_id = message.chat.id
    try:
        if str(UserRequests[message.chat.id]['request_category']).lower() == 'комментарии по рейтингу' or \
                str(UserRequests[message.chat.id]['request_category']).lower() == 'комментарии по тональности':
            send_str = 'Ваш выбор: ' + \
                       '\nЗапрос - ' + UserRequests[chat_id]['request_category'] + \
                       '\nТональность - ' + UserRequests[chat_id]['tonality'] + \
                       '\nПлатформа - ' + UserRequests[chat_id]['platform'] + \
                       '\nДата начала - ' + UserRequests[chat_id]['start_date'] + \
                       '\nДата конца - ' + UserRequests[chat_id]['end_date'] + \
                       '\nБренд - ' + UserRequests[chat_id]['brand'] + \
                       '\nФормат - ' + UserRequests[chat_id]['format']
        else:
            send_str = 'Ваш выбор: ' + \
                       '\nЗапрос - ' + UserRequests[chat_id]['request_category'] + \
                       '\nПлатформа - ' + UserRequests[chat_id]['platform'] + \
                       '\nДата начала - ' + UserRequests[chat_id]['start_date'] + \
                       '\nДата конца - ' + UserRequests[chat_id]['end_date'] + \
                       '\nБренд - ' + UserRequests[chat_id]['brand']
        bot.send_message(message.chat.id, send_str)
        markup = telebot.types.InlineKeyboardMarkup()
        markup.add(telebot.types.InlineKeyboardButton(text='Да', callback_data='Yes'))
        markup.add(telebot.types.InlineKeyboardButton(text='Нет', callback_data='No'))
        bot.send_message(message.chat.id, text='Все ли верно?', reply_markup=markup)
    except Exception as e:
        bot.send_message(message.chat.id, 'Пожалуйста, начните с начала и выберите категорию из меню')


def database_insert(chat_id):
    now = datetime.datetime.now()
    with closing(psycopg2.connect(dbname=PG_NAME, user=PG_USER, password=PG_PASSWORD, host=PG_HOST)) as conn:
        conn.autocommit = True
        with conn.cursor() as cursor:
            cursor.execute('INSERT INTO public."Users" VALUES ('+ str(chat_id)+
                ", '" + str(now)+ "', '" +str(UserRequests[chat_id]['platform'])+"', '" +
                str(UserRequests[chat_id]['start_date'])+"', '" +
                str(UserRequests[chat_id]['end_date'])+"', '" +str(UserRequests[chat_id]['brand'])+ "','"
                +str(UserRequests[chat_id]['request_category'])+"');")
    get_comments(chat_id)



def get_comments(chat_id):
    if str(UserRequests[chat_id]['platform']) == 'Wildberries':
        platform = 'wb'
    elif str(UserRequests[chat_id]['platform']) == 'Pet Shop':
        platform = 'petshop'
    elif str(UserRequests[chat_id]['platform']) == 'Мир корма':
        platform = 'mirkorma'
    request_str = 'http://backend-flask:5000/api/' + str(platform) \
                  + '/' + str(UserRequests[chat_id]['brand']) + '?date_start=' + \
                str(UserRequests[chat_id]['start_date']) + '&date_end=' + str(UserRequests[chat_id]['end_date'])
    try:
        r = requests.get(request_str, timeout=1000)
        #print(r.json())
        if r.status_code == 200:
            if r.json():
                choose_way_send(chat_id, r.json())
            else:
                bot.send_message(chat_id, 'Не удалось найти комментарии по вашему запросу. Проверьте запрос и попробуйте снова')
        else:
            raise Exception('Got error code from request - some problems in backend!')
    except Exception as e:
        print(e)
        bot.send_message(chat_id, 'Кажется, произошла ошибка. Проверьте запрос и попробуйте снова')
    # http://backend-flask:5000/api/petshop/whiskas?date_start=2023-05-01&date_end=2023-06-01
    #print(r.json)

#  'Общее облако слов' or call.data == 'Облако позитивных комментариев' or call.data == 'Облако негативных комментариев'

def choose_way_send(chat_id, dict_comments):
    if str(UserRequests[chat_id]['request_category']).lower() == 'общее облако слов':
        get_words_cloud_picture(chat_id, bot, dict_comments, str(UserRequests[chat_id]['brand'])  )

    elif str(UserRequests[chat_id]['request_category']).lower() == 'облако позитивных комментариев':
        coments_positive = get_one_tonality_comments(dict_comments, 0)
        get_words_cloud_picture(chat_id, bot, coments_positive, str(UserRequests[chat_id]['brand'])  )

    elif str(UserRequests[chat_id]['request_category']).lower() == 'облако негативных комментариев':
        coments_negative = get_one_tonality_comments(dict_comments, 1)
        get_words_cloud_picture(chat_id, bot, coments_negative, str(UserRequests[chat_id]['brand'])  )

    elif str(UserRequests[chat_id]['request_category']).lower() ==  'комментарии по тональности':
        if str(UserRequests[chat_id]['tonality']).lower() == 'негативные комментарии':
            choose_way_with_tonality(chat_id, dict_comments, 1)
        elif str(UserRequests[chat_id]['tonality']).lower() == 'позитивные комментарии':
            choose_way_with_tonality(chat_id, dict_comments, 0)

    elif str(UserRequests[chat_id]['request_category']).lower() ==  'комментарии по рейтингу':
        if str(UserRequests[chat_id]['tonality']).lower() == 'негативные комментарии':
            choose_way_with_rating(chat_id, dict_comments, 1)
        elif str(UserRequests[chat_id]['tonality']).lower() == 'позитивные комментарии':
            choose_way_with_rating(chat_id, dict_comments, 0)

    elif str(UserRequests[chat_id]['request_category']).lower() == 'статистика':
        get_comments_statistics(chat_id, dict_comments)




def choose_way_with_tonality(chat_id, dict_comments, tonality):
    coments_tonality = get_one_tonality_comments(dict_comments, tonality)
    if coments_tonality:
        if UserRequests[chat_id]['format'] == 'Word':
            send_word(chat_id, coments_tonality, 0)
        elif UserRequests[chat_id]['format'] == 'JSON':
            send_json(chat_id, coments_tonality)
    else:
        bot.send_message(chat_id, 'Не удалось найти комментарии с такой тональностью')

def choose_way_with_rating(chat_id, dict_comments, tonality):
    coments_rating = get_rating_comments(dict_comments, tonality)
    if coments_rating:
        if UserRequests[chat_id]['format'] == 'Word':
            send_word(chat_id, coments_rating, 1)
        elif UserRequests[chat_id]['format'] == 'JSON':
            send_json(chat_id, coments_rating)
    else:
        bot.send_message(chat_id, 'Не удалось найти комментарии с таким рейтингом')


def get_rating_comments(dict_comments, tonality):
    coments_good = []
    if tonality == 0:
        rating_lower = 4
        rating_upper = 5
    elif tonality == 1:
        rating_lower = 1
        rating_upper = 3
    for row in dict_comments:
        comments = []
        for com in row['comments']:
            if rating_lower <= com['rating'] <= rating_upper:
                comments.append(com)
        if comments:
            coments_good.append({'name': row['name'], 'url': row['url'],  'comments': comments})
    return coments_good




def get_one_tonality_comments(dict_comments, tonality):
    coments_good = []
    for row in dict_comments:
        comments = []
        for com in row['comments']:
            if com['tonality'] == tonality:  # positive
                comments.append(com)
        if comments:
            coments_good.append({'name': row['name'], 'url': row['url'],  'comments': comments})
    return coments_good



def get_comments_statistics(chat_id, dict_response):
    goods_number = len(dict_response)
    coments_good = []
    coments_bad = []
    coments_good_rating = []
    coments_bad_rating = []
    for row in dict_response:
        # comments = []
        for com in row['comments']:
            if com['tonality'] == 1:  # negative
                coments_bad.append(com)
            elif com['tonality'] == 0:  # positive
                coments_good.append(com)
            if com['rating'] >= 4:  # positive rating
                coments_good_rating.append(com)
            elif com['rating'] <= 3:  # negative rating
                coments_bad_rating.append(com)
    positive_num = len(coments_good)
    negative_num = len(coments_bad)
    positive_rat_num = len(coments_good_rating)
    negative_rat_num = len(coments_bad_rating)
    all_num = positive_num + negative_num
    send_statistics_str = 'Количество товаров бренда, на которые оставляли отзывы: ' + str(goods_number) + '\n\n'
    send_statistics_str += 'Количество комментариев: ' + str(all_num) + '\n\n'
    send_statistics_str += 'Количество положительных комментариев: ' + str(positive_num) + '\n\n'
    send_statistics_str += 'Количество комментариев с негативом: ' + str(negative_num) + '\n\n'
    send_statistics_str += 'Количество комментариев с рейтингом 4 или 5: ' + str(positive_rat_num) + '\n\n'
    send_statistics_str += 'Количество комментариев с рейтингом 3 и ниже: ' + str(negative_rat_num)
    bot.send_message(chat_id, send_statistics_str)






def send_word(chat_id, dict_comments, is_rating):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    for row in dict_comments:
        doc.add_paragraph(style='Normal').add_run(row['name']).bold = True
        doc.add_paragraph("Ссылка на товар: " + row['url'] + "\n", style='Normal')
        for com in row['comments']:
            if is_rating:
                doc.add_paragraph( (com['date']['$date'].replace('T', ' ').replace('Z', '') + " " + com['comment'] + "\n" +
                     "Рейтинг: " + str(com['rating']) + "\n"), style='Normal')
            else:
                doc.add_paragraph((com['date']['$date'].replace('T', ' ').replace('Z', '') + " " + com['comment'] + "\n"), style='Normal')

    doc_name = UserRequests[chat_id]['tonality'].partition(' ')[0]  \
                + ' ' + str(UserRequests[chat_id]['platform'])  \
                + ' ' + str(UserRequests[chat_id]['brand']) + ' от ' + \
                str(UserRequests[chat_id]['start_date']) + ' до ' + str(UserRequests[chat_id]['end_date']) + ".docx"
    doc.save(doc_name)
    with open(doc_name, "rb") as f:
        bot.send_document(chat_id, f)
    os.remove(doc_name)




def send_json(chat_id, dict_comments):
    doc_name = str(UserRequests[chat_id]['platform']) \
               + ' ' + str(UserRequests[chat_id]['brand']) + ' от ' + \
               str(UserRequests[chat_id]['start_date']) + ' до ' + str(UserRequests[chat_id]['end_date']) + ".txt"
    with open(doc_name, "w") as f:
        f.write(str(dict_comments))
    with open(doc_name, "r") as f:
        bot.send_document(chat_id, f)
    os.remove(doc_name)



@bot.callback_query_handler(func=lambda call: True)
def query_handler(call):
    markup = telebot.types.InlineKeyboardMarkup()
    markup.add(telebot.types.InlineKeyboardButton(text='Последняя неделя', callback_data='last_week'))
    markup.add(telebot.types.InlineKeyboardButton(text='Последний месяц', callback_data='last_month'))
    markup.add(telebot.types.InlineKeyboardButton(text='Последний год', callback_data='last_year'))
    markup.add(telebot.types.InlineKeyboardButton(text='Свой вариант', callback_data='other_interval'))
    
    #bot.answer_callback_query(callback_query_id=call.id, text='Спасибо за честный ответ!')
    answer = ''
    if call.data == 'Wildberries' or call.data == 'Pet Shop' or call.data == 'Мир корма':
        platform = call.data
        try:
            UserRequests[call.message.chat.id]['platform'] = platform
            answer = 'Супер!'
            bot.send_message(call.message.chat.id, text="Выберите даты, за которые будут выданы комментарии", reply_markup=markup)
        except:
            bot.send_message(call.message.chat.id, 'Пожалуйста, начните с начала и выберите категорию из меню')


    elif call.data == 'Позитивные комментарии' or call.data =='Негативные комментарии':
        UserRequests[call.message.chat.id]['tonality'] = call.data
        platform_choosing(call.message)



    elif call.data == 'Общее облако слов' or call.data == 'Облако позитивных комментариев' or call.data == 'Облако негативных комментариев':
        request_category = call.data
        UserRequests[call.message.chat.id] = {"id": call.message.chat.id, "request_category": request_category,
                                 "platform": None, "start_date": None, "end_date": None, "brand": None}
        platform_choosing(call.message)



    elif call.data == 'last_week':
        start_date,end_date = get_standart_date_interval(1)
        dates_interval_set(call.message, start_date, end_date)
    elif call.data == 'last_month':
        start_date, end_date = get_standart_date_interval(4)
        dates_interval_set(call.message, start_date, end_date)
    elif call.data == 'last_year':
        start_date, end_date = get_standart_date_interval(52)
        dates_interval_set(call.message, start_date, end_date)
    elif call.data == 'other_interval':
        dates_interval_choosing(call.message)


    elif call.data == 'Yes':
        bot.send_message(call.message.chat.id, "Отлично! Подождите немного пока расчитывается результат")
        #bot.send_video(call.message.chat.id, 'https://i.pinimg.com/originals/26/2e/3a/262e3a9f491af6c30ed91e4263bd19b8.gif', None, 'Text')
        #https://usagif.com/wp-content/uploads/loading-96.gif
        bot.send_sticker(call.message.chat.id, 'CAACAgIAAxkBAAEI5bJkWJgWRYc2MSyyrtyzPrfLyMrvnAACUgADr8ZRGgSvecXtKHqOLwQ')
        markup = telebot.types.InlineKeyboardMarkup()
        database_insert(call.message.chat.id)

    elif call.data == 'No':
        bot.send_message(call.message.chat.id, "Нам очень жаль! Попробуйте еще раз")
        platform_choosing(call.message)


    elif call.data == 'Word' or call.data == 'JSON':
        UserRequests[call.message.chat.id]['format'] = call.data
        check_info(call.message)



    bot.edit_message_reply_markup(call.message.chat.id, call.message.message_id)


bot.polling()



